import logging
import os

from docx import Document
from docx.shared import Inches as DocxInches


def generate_evidence():
    """Gera um arquivo de evidência com base nos screenshots capturados."""
    model_path = "modelos de evidencias/1.evidencia.docx"
    screenshots_dir = "reports/screenshots"
    output_path = "reports/evidencias/Evidencia_Atualizada.docx"

    # Garante que a pasta de screenshots exista
    if not os.path.exists(screenshots_dir):
        logging.warning(f"Diretório de screenshots não encontrado: {screenshots_dir}")
        return None

    # Carrega o modelo de evidência
    document = Document(model_path)

    # Adiciona os passos numerados e seus prints ao documento
    step_number = 1
    for file_name in sorted(os.listdir(screenshots_dir)):
        if file_name.endswith(".png"):
            step_title = file_name.replace("_", " ").replace(".png", "").capitalize()
            try:
                # Adiciona o número e o nome do passo em negrito
                paragraph = document.add_paragraph()
                run = paragraph.add_run(f"{step_number}. {step_title}")
                run.bold = True
                step_number += 1
            except Exception as e:
                logging.warning(
                    f"Erro ao adicionar o título do passo '{step_title}': {e}"
                )
            try:
                # Adiciona a imagem correspondente ao passo e centraliza
                paragraph = document.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(
                    os.path.join(screenshots_dir, file_name), width=DocxInches(5)
                )
                paragraph.alignment = 1  # Centraliza a imagem
            except Exception as e:
                logging.warning(f"Erro ao adicionar a imagem '{file_name}': {e}")

    # Salva o arquivo de evidência
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    document.save(output_path)
    logging.info(f"Evidência gerada: {output_path}")
    return output_path
