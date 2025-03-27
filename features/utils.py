import os
import logging
from datetime import datetime  # Corrigido para evitar redundância
from docx import Document
from docx2pdf import convert

def gerar_documento_evidencia(nome_teste, sucesso=True, erros=None):
    """Gera um documento de evidência ou bug com base no modelo fornecido."""
    try:
        data_teste = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        modelo = "modelos de evidencias/1.evidencia.docx" if sucesso else "modelos de evidencias/2.bug.docx"
        nome_arquivo = f"reports/evidencias/{'Evidência' if sucesso else 'Bug'}_{data_teste}.docx"

        doc = Document(modelo)
        doc.add_paragraph(f"Teste: {nome_teste}")
        doc.add_paragraph(f"Data do Teste: {data_teste}")

        if not sucesso and erros:
            doc.add_paragraph("Erros encontrados:")
            for erro in erros:
                doc.add_paragraph(f"- {erro}")

        os.makedirs(os.path.dirname(nome_arquivo), exist_ok=True)
        doc.save(nome_arquivo)
        logging.info(f"Documento gerado: {nome_arquivo}")
    except Exception as e:
        logging.error(f"Erro ao gerar documento de evidência: {e}")
        raise

def gerar_resumo_testes(total_testes, testes_sucesso, testes_falha):
    """
    Gera um documento de resumo com métricas dos testes realizados.
    """
    modelo = "modelos de evidencias/3.Resumo.docx"
    data_teste = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")  # Corrigido para usar datetime.now()
    nome_arquivo = f"reports/evidencias/Resumo_{data_teste}.docx"

    doc = Document(modelo)
    doc.add_paragraph(f"Data do Teste: {data_teste}")
    doc.add_paragraph(f"Total de Testes: {total_testes}")
    doc.add_paragraph(f"Testes com Sucesso: {testes_sucesso}")
    doc.add_paragraph(f"Testes com Falha: {testes_falha}")

    # Garante que a pasta "reports/evidencias" exista
    os.makedirs(os.path.dirname(nome_arquivo), exist_ok=True)

    doc.save(nome_arquivo)
    print(f"Resumo gerado: {nome_arquivo}")
