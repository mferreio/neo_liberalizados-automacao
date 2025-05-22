from pages.diretriz_irec_pages import DiretrizIrecPage
import base64
import io
import os
import traceback
import shutil
import stat
import subprocess
from docx import Document
from datetime import datetime
from time import sleep
import docx.shared
import matplotlib.pyplot as plt
import pyautogui
from pages.login_page import LoginPage, LoginPageLocators
from pages.perfil_de_acesso_pages import PerfilDeAcessoPage
from selenium import webdriver
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait

from credentials import LOGIN_EMAIL, LOGIN_PASSWORD, LOGIN_USUARIO


def login(context):
    """Realiza o login no sistema."""
    try:
        context.login_page.navegar_para_pagina_de_login()
        context.login_page.clicar_botao_entrar()
        context.login_page.enter_email(context.login_email)
        context.login_page.click_next_button()
        sleep(8)
        pyautogui.write(LOGIN_USUARIO.upper())
        pyautogui.press("tab")
        sleep(1)
        pyautogui.write(LOGIN_PASSWORD)
        pyautogui.press("tab")
        sleep(1)
        pyautogui.press("enter")
        sleep(10)

        # Aguarda a transição para a próxima página
        WebDriverWait(context.driver, 15).until(
            EC.url_contains("https://diretrizes.dev.neoenergia.net/")
        )
        # logging removido
    except TimeoutException as e:
        # logging removido
        context.driver.save_screenshot("reports/screenshots/timeout_exception.png")
        raise


def before_all(context):
    """Configura o ambiente antes de todos os testes."""

    # Limpa e garante que as pastas necessárias existam (ANTES de iniciar o Selenium e contexto)
    import shutil
    shutil.rmtree("reports/allure-results", ignore_errors=True)
    shutil.rmtree("reports/allure-report", ignore_errors=True)
    os.makedirs("reports/screenshots", exist_ok=True)
    os.makedirs("reports/evidencias", exist_ok=True)
    os.makedirs("reports/allure-results", exist_ok=True)

    fixed_port = 8080  # Porta fixa para o Selenium

    chrome_options = Options()
    chrome_options.add_argument("--start-maximized")
    chrome_options.add_argument("--incognito")
    chrome_options.add_argument("--ignore-certificate-errors")
    chrome_options.add_argument(
        f"--remote-debugging-port={fixed_port}"
    )  # Configura a porta
    context.driver = webdriver.Chrome(options=chrome_options)
    context.driver.delete_all_cookies()
    context.login_page = LoginPage(context.driver)
    context.login_email = LOGIN_EMAIL
    context.passed_steps = []
    context.failed_steps = []  # Lista para armazenar os erros capturados
    context.passed_scenarios = []
    context.failed_scenarios = []
    context.start_time = datetime.now()

    context.report_output_path = os.path.join(
        os.getcwd(), "docs"
    )  # Define o diretório de saída como 'docs'
    if not os.path.exists(context.report_output_path):
        os.makedirs(context.report_output_path)  # Cria o diretório se não existir
    context.github_pages_url = (
        "https://mferreio.github.io/neo_liberalizados-automacao/"  # URL do GitHub Pages
    )
    context.github_pages_branch = (
        "gh-pages"  # Define a branch usada para o GitHub Pages
    )

    context.perfil_de_acesso_pages = PerfilDeAcessoPage(context.driver)

    # Realiza o login apenas uma vez antes de todos os testes
    try:
        login(context)
    except Exception as e:
        # logging removido
        raise


def before_feature(context, feature):
    """Executa ações antes de cada feature."""
    # logging removido

    # Verifica se a feature é '09_perfil_de_acesso_nao_logado.feature'
    if "09_perfil_de_acesso_nao_logado.feature" in feature.filename:
        try:
            # Encerra o navegador para garantir que não há sessão logada
            if hasattr(context, "driver") and context.driver is not None:
                try:
                    context.driver.quit()
                except Exception:
                    pass
                context.driver = None
            # O driver será criado no primeiro step dessa feature, garantindo navegador limpo
        except Exception as e:
            raise
    else:
        if hasattr(context, "driver") and context.driver is not None:
            context.driver.get("https://diretrizes.dev.neoenergia.net/")


def before_scenario(context, scenario):
    # Inicializa lista de evidências dos passos para o cenário
    context.evidencias_passos = []
    # Agrupamento por feature
    if not hasattr(context, 'evidencias_features'):
        context.evidencias_features = {}
    feature_path = scenario.feature.filename
    if feature_path not in context.evidencias_features:
        context.evidencias_features[feature_path] = []

    # --- INÍCIO DO AJUSTE: cria driver limpo para feature de não logado ---
    if (
        "09_perfil_de_acesso_nao_logado.feature" in feature_path
        and (not hasattr(context, "driver") or context.driver is None)
    ):
        chrome_options = Options()
        chrome_options.add_argument("--start-maximized")
        chrome_options.add_argument("--incognito")
        chrome_options.add_argument("--ignore-certificate-errors")
        chrome_options.add_argument("--remote-debugging-port=8080")
        context.driver = webdriver.Chrome(options=chrome_options)
        context.driver.delete_all_cookies()
        # Não executa login!
    # --- FIM DO AJUSTE ---

def after_step(context, step):
    """Captura evidência (screenshot) e status de todos os passos, independente de sucesso ou falha."""
    import os
    from datetime import datetime
    screenshot_path = os.path.join('reports', 'screenshots', f"step_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}.png")
    try:
        if hasattr(context, 'driver'):
            context.driver.save_screenshot(screenshot_path)
    except Exception:
        screenshot_path = None
    status = 'Sucesso' if step.status == 'passed' else 'Falha'
    error = str(step.exception) if step.status == 'failed' and step.exception else None
    if hasattr(context, 'evidencias_passos'):
        context.evidencias_passos.append({
            'step': step.name,
            'status': status,
            'screenshot': screenshot_path,
            'error': error
        })
    # Exclui o print após alimentar a evidência
    if screenshot_path and os.path.exists(screenshot_path):
        try:
            os.remove(screenshot_path)
        except Exception:
            pass
    # Inicializa o page object DiretrizIrecPage para todos os cenários que envolvem diretriz I-REC
    if hasattr(context, "driver"):
        context.diretriz_irec_page = DiretrizIrecPage(context.driver)
    """Executa ações antes de cada cenário."""
    # logging removido
    context.start_time_scenario = datetime.now()  # Registra o início do cenário


def after_scenario(context, scenario):
    # Ao final de cada cenário, armazena as evidências dos passos no agrupamento da feature
    feature_path = scenario.feature.filename
    if hasattr(context, 'evidencias_features'):
        context.evidencias_features[feature_path].append({
            'cenario': scenario.name,
            'status': scenario.status,
            'passos': list(context.evidencias_passos)
        })
    """Executa ações após cada cenário."""
    end_time_scenario = datetime.now()
    execution_time = end_time_scenario - context.start_time_scenario
    # logging removido
    if scenario.status == "failed":
        # logging removido
        context.failed_scenarios.append(scenario.name)
    else:
        context.passed_scenarios.append(scenario.name)


def esperar_e_executar(context, locator, metodo, *args):
    """Espera por um elemento clicável, executa uma ação e registra evidência do passo."""
    from time import sleep
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.support.ui import WebDriverWait
    import os
    import traceback
    try:
        WebDriverWait(context.driver, 20).until(EC.element_to_be_clickable(locator))
        metodo(*args)
        status = 'Sucesso'
        error = None
    except Exception as e:
        status = 'Falha'
        error = str(e) + '\n' + traceback.format_exc()
        raise
    finally:
        sleep(1)
        # Captura screenshot e registra passo
        screenshot_path = os.path.join('reports', 'screenshots', f"step_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}.png")
        try:
            context.driver.save_screenshot(screenshot_path)
        except Exception:
            screenshot_path = None
        # Descobre nome do passo
        import inspect
        frame = inspect.currentframe().f_back
        step_name = frame.f_code.co_name if frame else 'passo_desconhecido'
        if hasattr(context, 'evidencias_passos'):
            context.evidencias_passos.append({
                'step': step_name,
                'status': status,
                'screenshot': screenshot_path,
                'error': error
            })


def gerar_documento_evidencia(nome_teste, sucesso=True, erros=None):
    """
    Gera um documento de evidência ou bug com base no modelo fornecido.
    Sempre utiliza o template 1.evidencia.docx.
    """
    import os
    from docx import Document

    data_teste = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    # Sempre usa o template 1.evidencia.docx
    modelo = os.path.abspath(os.path.join(os.getcwd(), "modelos de evidencias", "1.evidencia.docx"))
    if sucesso:
        nome_arquivo = f"reports/evidencias/Evidência_{data_teste}.docx"
    else:
        nome_arquivo = f"reports/evidencias/Bug_{data_teste}.docx"

    doc = Document(modelo)
    doc.add_paragraph(f"Teste: {nome_teste}")
    doc.add_paragraph(f"Data do Teste: {data_teste}")

    if not sucesso and erros:
        doc.add_paragraph("Erros encontrados:")
        for erro in erros:
            doc.add_paragraph(f"- {erro}")

    # Garante que a pasta "reports/evidencias" exista
    os.makedirs(os.path.dirname(nome_arquivo), exist_ok=True)

    doc.save(nome_arquivo)
    # logging removido
    return nome_arquivo


def gerar_resumo_testes(total_testes, testes_sucesso, testes_falha):
    """
    Gera um documento de resumo com métricas dos testes realizados.
    """
    import os

    from docx import Document

    modelo = "modelos de evidencias/3.Resumo.docx"
    data_teste = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    nome_arquivo = f"reports/evidencias/Resumo_{data_teste}.docx"

    doc = Document(modelo)
    doc.add_paragraph(f"Data do Teste: {data_teste}")
    doc.add_paragraph(f"Total de Testes: {total_testes}")
    doc.add_paragraph(f"Testes com Sucesso: {testes_sucesso}")
    doc.add_paragraph(f"Testes com Falha: {testes_falha}")

    # Garante que a pasta "reports/evidencias" exista
    os.makedirs(os.path.dirname(nome_arquivo), exist_ok=True)

    doc.save(nome_arquivo)
    # logging removido
    return nome_arquivo


def reset_permissions(directory):
    """Redefine permissões de um diretório e seus arquivos."""
    for root, dirs, files in os.walk(directory):
        for dir_name in dirs:
            os.chmod(os.path.join(root, dir_name), stat.S_IRWXU)
        for file_name in files:
            os.chmod(os.path.join(root, file_name), stat.S_IRWXU)


def gerar_grafico_percentual(total, falhas, titulo):
    """Gera um gráfico de pizza com o percentual de falhas."""
    sucesso = total - falhas
    labels = ["Sucesso", "Falhas"]
    sizes = [sucesso, falhas]
    colors = ["#4CAF50", "#F44336"]
    explode = (0, 0.1)  # Destaque para falhas

    # Evita divisão por zero ao gerar o gráfico
    if total == 0:
        sizes = [1]  # Exibe 100% como "Nenhum dado"
        labels = ["Nenhum dado"]
        colors = ["#B0BEC5"]  # Cor neutra para ausência de dados
        explode = (0,)

    fig, ax = plt.subplots(figsize=(4, 4))  # Reduz o tamanho do gráfico
    ax.pie(
        sizes,
        explode=explode,
        labels=labels,
        colors=colors,
        autopct="%1.1f%%",
        startangle=90,
    )
    ax.axis("equal")  # Garante que o gráfico seja um círculo
    ax.set_title(titulo, fontsize=10)

    # Salva o gráfico em memória
    buffer = io.BytesIO()
    plt.savefig(buffer, format="png", bbox_inches="tight")
    buffer.seek(0)
    img_base64 = base64.b64encode(buffer.read()).decode("utf-8")
    buffer.close()
    plt.close(fig)
    return img_base64


def gerar_grafico_percentual_completo(sucesso, falhas, ignorados, titulo):
    """Gera um gráfico de pizza com os percentuais de sucesso, falhas e ignorados."""
    total = sucesso + falhas + ignorados
    labels = ["Sucesso", "Falhas", "Ignorados"]
    sizes = [sucesso, falhas, ignorados]
    colors = ["#4CAF50", "#F44336", "#FFC107"]
    explode = (0, 0.1, 0)  # Destaque para falhas

    # Evita divisão por zero ao gerar o gráfico
    if total == 0:
        sizes = [1]  # Exibe 100% como "Nenhum dado"
        labels = ["Nenhum dado"]
        colors = ["#B0BEC5"]  # Cor neutra para ausência de dados
        explode = (0,)

    fig, ax = plt.subplots(figsize=(4, 4))  # Reduz o tamanho do gráfico
    ax.pie(
        sizes,
        explode=explode,
        labels=labels,
        colors=colors,
        autopct="%1.1%%",
        startangle=90,
    )
    ax.axis("equal")  # Garante que o gráfico seja um círculo
    ax.set_title(titulo, fontsize=10)

    # Salva o gráfico em memória
    buffer = io.BytesIO()
    plt.savefig(buffer, format="png", bbox_inches="tight")
    buffer.seek(0)
    img_base64 = base64.b64encode(buffer.read()).decode("utf-8")
    buffer.close()
    plt.close(fig)
    return img_base64

def after_all(context):
    """Finaliza o ambiente após todos os testes."""
    try:
        if hasattr(context, "driver"):
            context.driver.quit()

        # Geração de evidências por feature (mantido)
        print('DEBUG: after_all chamado')
        print('DEBUG: evidencias_features:', getattr(context, 'evidencias_features', None))
        if hasattr(context, 'evidencias_features') and context.evidencias_features:
            for feature_path, cenarios in context.evidencias_features.items():
                print(f"DEBUG: Gerando evidência para feature: {feature_path}")
                feature_nome = os.path.splitext(os.path.basename(feature_path))[0]
                data_execucao = datetime.now().strftime('%d-%m-%Y')
                hora_execucao = datetime.now().strftime('%H-%M-%S')
                responsavel = os.getenv('USUARIO_TESTE', 'Automação')
                dispositivo = 'Chrome'
                versao = ''
                info = {
                    'projeto': 'Liberalizados',
                    'frente': 'Web',
                    'responsavel': responsavel,
                    'distribuidora': 'NE e SE',
                    'data': data_execucao,
                    'produto': 'Aplicação Base',
                    'resultado_esperado': 'Execução dos passos dos cenários',
                    'versao': versao,
                    'dispositivo': dispositivo
                }

                # Sempre usa o template 1.evidencia.docx
                modelo = os.path.abspath(os.path.join(os.getcwd(), "modelos de evidencias", "1.evidencia.docx"))
                print(f"DEBUG: Procurando template em: {modelo}")
                if not os.path.exists(modelo):
                    print(f"ERRO: Template não encontrado em: {modelo}")
                    continue
                doc = Document(modelo)
                doc.add_paragraph(f"Projeto: {info['projeto']}")
                doc.add_paragraph(f"Frente: {info['frente']}")
                doc.add_paragraph(f"Responsável: {info['responsavel']}")
                doc.add_paragraph(f"Distribuidora: {info['distribuidora']}")
                doc.add_paragraph(f"Data Execução: {data_execucao}")
                doc.add_paragraph(f"Hora Execução: {hora_execucao}")
                doc.add_paragraph(f"Produto: {info['produto']}")
                doc.add_paragraph(f"Resultado Esperado: {info['resultado_esperado']}")
                doc.add_paragraph(f"Versão: {info['versao']}")
                doc.add_paragraph(f"Dispositivo: {info['dispositivo']}")
                doc.add_paragraph(f"Feature: {feature_nome}")

                for cenario in cenarios:
                    doc.add_paragraph(f"\nCenário: {cenario['cenario']}")
                    doc.add_paragraph(f"Status: {'Sucesso' if cenario['status']=='passed' else 'Falha'}")
                    for idx, passo in enumerate(cenario['passos'], 1):
                        doc.add_paragraph(f"{idx} - {passo['step']}")
                        if passo.get('screenshot') and os.path.exists(passo['screenshot']):
                            try:
                                doc.add_picture(passo['screenshot'], width=docx.shared.Inches(4))
                            except Exception as e:
                                doc.add_paragraph(f"[Erro ao inserir imagem: {e}]")
                        doc.add_paragraph(f"Status: {passo['status']}")
                        if passo.get('error'):
                            doc.add_paragraph(f"Erro: {passo['error']}")

                # Caminho absoluto para a pasta evidencias na raiz do repositório
                pasta = os.path.abspath(os.path.join(os.getcwd(), "reports", "evidencias"))
                os.makedirs(pasta, exist_ok=True)
                nome_arquivo = os.path.join(
                    pasta,
                    f"Evidencia_{feature_nome}_{data_execucao}_{hora_execucao}.docx"
                )
                print(f"DEBUG: Salvando arquivo de evidência em: {nome_arquivo}")
                try:
                    doc.save(nome_arquivo)
                    print(f"Arquivo de evidência gerado: {nome_arquivo}")
                except Exception as e:
                    print(f"Erro ao salvar arquivo de evidência {nome_arquivo}: {e}")

        # Gera o relatório do Allure ao final dos testes
        allure_report_generated = False
        try:
            allure_executable = r"C:\allure\bin\allure.bat"
            allure_results_dir = "reports/allure-results"
            allure_report_dir = "reports/allure-report"
            subprocess.run(
                [
                    allure_executable,
                    "generate",
                    allure_results_dir,
                    "-o",
                    allure_report_dir,
                    "--clean",
                ],
                check=True,
            )
            print(f"Relatório Allure gerado em: {allure_report_dir}")
            allure_report_generated = True
        except Exception as e:
            print(f"Erro ao gerar o relatório Allure: {e}")

        # Copia o relatório Allure para docs/allure-report-<data>-<N>/ e faz commit/push para main
        report_url = "Erro ao gerar link do relatório"
        if allure_report_generated and os.path.exists("reports/allure-report"):
            try:
                # Gera nome com data (DD-MM-AAAA) e contador incremental
                data_str = datetime.now().strftime('%d-%m-%Y')
                docs_base = os.path.join('docs')
                # Busca todos os diretórios já existentes para a data
                existing = [
                    d for d in os.listdir(docs_base)
                    if d.startswith(f'allure-report-{data_str}')
                ]
                # Calcula o próximo índice
                idx = 1
                if existing:
                    nums = []
                    for d in existing:
                        parts = d.split('-')
                        if len(parts) >= 4 and parts[2] == data_str.split('-')[0] and parts[3] == data_str.split('-')[1] and parts[4] == data_str.split('-')[2]:
                            try:
                                nums.append(int(parts[-1]))
                            except Exception:
                                pass
                    if nums:
                        idx = max(nums) + 1
                report_dir_name = f'allure-report-{data_str}-{idx}'
                docs_report_path = os.path.join(docs_base, report_dir_name)
                if os.path.exists(docs_report_path):
                    shutil.rmtree(docs_report_path)
                shutil.copytree('reports/allure-report', docs_report_path)
                print(f'Relatório Allure copiado para: {docs_report_path}')

                # Commit e push para main (ignora hooks)
                from git import Repo
                repo = Repo(os.getcwd())
                repo.git.add(docs_report_path)
                repo.git.commit('-m', f'Adiciona relatório Allure {report_dir_name} em docs para histórico no GitHub Pages', '--no-verify')
                repo.git.push('origin', 'main')
                print('Relatório Allure enviado para a branch main (GitHub Pages).')

                # Monta o link único do relatório
                github_pages_base = "https://mferreio.github.io/neo_liberalizados-automacao/"
                report_url = f"{github_pages_base}{report_dir_name}/"
            except Exception as e:
                print(f"Erro ao copiar/commit/push do relatório Allure para docs/: {e}")
                report_url = "Erro ao gerar link do relatório"
        else:
            print("Relatório Allure não foi gerado, não será copiado para docs/.")

        # --- Envia o e-mail com o link único do relatório Allure ---
        try:
            # Salva as métricas em variáveis de ambiente para o script de e-mail
            os.environ['ALLURE_REPORT_URL'] = report_url
            os.environ['DATA_EXECUCAO'] = datetime.now().strftime('%d/%m/%Y %H:%M')
            os.environ['TOTAL_TESTES'] = str(len(getattr(context, 'passed_scenarios', [])) + len(getattr(context, 'failed_scenarios', [])))
            os.environ['TESTES_APROVADOS'] = str(len(getattr(context, 'passed_scenarios', [])))
            os.environ['TESTES_FALHADOS'] = str(len(getattr(context, 'failed_scenarios', [])))
            os.environ['TESTES_IGNORADOS'] = "0"  # Ajuste se necessário

            # Calcula o tempo de execução total
            end_time = datetime.now()
            if hasattr(context, 'start_time'):
                execution_time = end_time - context.start_time
                minutos, segundos = divmod(execution_time.total_seconds(), 60)
                tempo_execucao_str = f"{int(minutos)}m {int(segundos)}s"
            else:
                tempo_execucao_str = "-"
            os.environ['TEMPO_EXECUCAO'] = tempo_execucao_str

            subprocess.run(
                ["python", "send_allure_report_email.py"],
                check=True
            )
        except Exception as e:
            print(f"Erro ao chamar o envio de e-mail: {e}")

        # Calcula o tempo de execução
        end_time = datetime.now()
        execution_time = end_time - context.start_time
    except Exception as e:
        print("Erro ao gerar evidência:", e)
        traceback.print_exc()