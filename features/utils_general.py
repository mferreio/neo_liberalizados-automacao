import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.image import MIMEImage
import os
from docx import Document
from subprocess import run
import logging
from docx.shared import Inches as DocxInches
import matplotlib.pyplot as plt
import datetime

def send_email(subject, body, recipient_email, attachment_path=None, image_path=None):
    """Envia um e-mail com ou sem anexos."""
    try:
        sender_email = os.getenv("REMETENTE_DE_EMAIL")
        password = os.getenv("APP_PASSWORD")

        if not sender_email or not password:
            raise ValueError("As credenciais de e-mail não estão configuradas corretamente no arquivo .env.")

        msg = MIMEMultipart()
        msg['Subject'] = subject
        msg['From'] = sender_email
        msg['To'] = recipient_email

        # Corpo do e-mail
        msg.attach(MIMEText(body, 'html'))

        # Anexa imagem, se fornecida
        if image_path:
            with open(image_path, 'rb') as img:
                mime_img = MIMEImage(img.read())
                mime_img.add_header('Content-ID', '<image1>')
                msg.attach(mime_img)

        # Anexa arquivo, se fornecido
        if attachment_path:
            with open(attachment_path, 'rb') as attachment:
                part = MIMEText(attachment.read(), 'base64', 'utf-8')
                part.add_header('Content-Disposition', f'attachment; filename="{os.path.basename(attachment_path)}"')
                msg.attach(part)

        # Envia o e-mail
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(sender_email, password)
            server.send_message(msg)

        logging.info(f"E-mail enviado com sucesso para {recipient_email}")
    except Exception as e:
        logging.error(f"Erro ao enviar o e-mail: {e}")
        raise

def generate_evidence():
    """Gera um arquivo de evidência com base nos screenshots capturados."""
    model_path = 'modelos de evidencias/1.evidencia.docx'
    document = Document(model_path)
    screenshots_dir = 'reports/screenshots'

    # Garante que a pasta de screenshots exista
    if not os.path.exists(screenshots_dir):
        logging.warning(f"Diretório de screenshots não encontrado: {screenshots_dir}")
        return None

    # Adiciona os passos numerados e seus prints ao documento
    step_number = 1
    for file_name in sorted(os.listdir(screenshots_dir)):
        if file_name.endswith('.png'):
            step_title = file_name.replace("_", " ").replace(".png", "").capitalize()
            try:
                # Adiciona o número e o nome do passo em negrito
                paragraph = document.add_paragraph()
                run = paragraph.add_run(f"{step_number}. {step_title}")
                run.bold = True
                step_number += 1
            except Exception as e:
                logging.warning(f"Erro ao adicionar o título do passo '{step_title}': {e}")
            try:
                # Adiciona a imagem correspondente ao passo e centraliza
                paragraph = document.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(os.path.join(screenshots_dir, file_name), width=DocxInches(5))
                paragraph.alignment = 1  # Centraliza a imagem
            except Exception as e:
                logging.warning(f"Erro ao adicionar a imagem '{file_name}': {e}")
    
    output_path = 'reports/evidencias/Evidencia_Atualizada.docx'
    document.save(output_path)
    logging.info(f"Evidência gerada: {output_path}")
    return output_path
